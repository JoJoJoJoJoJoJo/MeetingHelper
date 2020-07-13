from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX

import os
import logging

from config import Config

_logger = logging.getLogger(__name__)


class WordGenerator:
    def __init__(self, name, contents, show_details=True):
        self.name = name
        self.contents = contents
        self.show_details = show_details
        self.filename = self.name.rsplit('.', maxsplit=1)[0] + ('（含详情）' if show_details else '') + '.docx'
        self.file_path = self.get_file_path(self.filename)
        try:
            self.clear_docx()
        except Exception as e:
            _logger.warning(e)

    def generate(self):
        # self.merge_contents()
        document = Document()

        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for content in self.contents:
            begin = content['bg']
            end = content['ed']
            speaker = content['speaker']
            text = content['onebest']
            paragh = document.add_paragraph('')
            if self.show_details:
                run = paragh.add_run('第{}秒至第{}秒，发言人{}号'.format(int(begin) / 1000, int(end) / 1000, speaker))
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_2
                paragh.add_run('\n')
            paragh.add_run(text)
        document.save(self.file_path)

    def merge_contents(self):
        """
        容易有坑，先放下
        :return:
        """
        index = 0
        keep_indexes = []
        while index < len(self.contents) - 1:
            content = self.contents[index]
            if content['onebest'][-1] not in ['。', '！', '？'] and content['speaker'] == self.contents[index + 1]['speaker']:
                # 如果不是一句话结尾，且归属于同一发言人，则向后合并
                next_content = self.contents[index + 1]
                next_content['bg'] = content['bg']
                next_content['onebest'] = content['onebest'] + next_content['onebest']
            else:
                keep_indexes.append(index)
            index += 1
        self.contents = [self.contents[i] for i in keep_indexes]

    def clear_docx(self):
        files = os.listdir(Config.WORD_SAVE_DIR)
        for name in filter(lambda f: f.endswith('.docx'), files):
            os.remove(name)

    @classmethod
    def get_file_path(cls, filename):
        return os.path.join(Config.WORD_SAVE_DIR, filename)
